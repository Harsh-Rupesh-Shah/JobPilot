"""
backend/tools/file_export.py
Tools to export agent outputs to DOCX and PDF.
"""
import logging
import os
from pathlib import Path
from backend.agents.state import ApplicationState
from docx import Document

logger = logging.getLogger(__name__)

async def export_outputs(state: ApplicationState) -> None:
    """
    Exports the tailored_resume and cover_letter to DOCX files.
    """
    run_id = state.get("run_id")
    if not run_id:
        return
        
    output_dir = Path("uploads/outputs") / run_id
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Export resume
    resume_content = state.get("tailored_resume", "")
    if resume_content:
        doc = Document()
        for line in resume_content.split("\\n"):
            doc.add_paragraph(line)
        resume_path = output_dir / "tailored_resume.docx"
        doc.save(resume_path)
        logger.info("Exported resume to %s", resume_path)
        
    # Export cover letter
    cover_letter_content = state.get("cover_letter", "")
    if cover_letter_content:
        doc = Document()
        for line in cover_letter_content.split("\\n"):
            doc.add_paragraph(line)
        cl_path = output_dir / "cover_letter.docx"
        doc.save(cl_path)
        logger.info("Exported cover letter to %s", cl_path)
